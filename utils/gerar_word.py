from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def gerar_word_orcamento(dados):
    buffer = BytesIO()
    doc = Document()
    
    # Título
    title = doc.add_heading('Orçamento Auto Mecânica Piau', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do Cliente
    doc.add_heading('Informações do Cliente', level=1)
    doc.add_paragraph(f"Cliente: {dados['cliente']['nome']}")
    doc.add_paragraph(f"Telefone: {dados['cliente']['telefone']}")
    doc.add_paragraph(f"Veículo: {dados['cliente']['veiculo']} - Placa: {dados['cliente']['placa']}")
    doc.add_paragraph(f"Data: {dados['cliente']['data']}")
    
    # Peças
    doc.add_heading('Peças e Componentes', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Peça'
    hdr_cells[1].text = 'Qtd'
    hdr_cells[2].text = 'V. Unit'
    hdr_cells[3].text = 'V. Total'
    
    for peca in dados['pecas']:
        row_cells = table.add_row().cells
        row_cells[0].text = peca['nome']
        row_cells[1].text = str(peca['quantidade'])
        row_cells[2].text = f"R$ {peca['valor_unitario']:.2f}"
        row_cells[3].text = f"R$ {peca['valor_total']:.2f}"
    
    # Mão de Obra
    doc.add_heading('Serviços e Mão de Obra', level=1)
    doc.add_paragraph(f"Descrição: {dados['mao_de_obra']['descricao']}")
    doc.add_paragraph(f"Valor da Mão de Obra: R$ {dados['mao_de_obra']['valor']:.2f}")
    
    # Total Geral
    total_para = doc.add_paragraph()
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = total_para.add_run(f"VALOR TOTAL: R$ {dados['total_geral']:.2f}")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer
